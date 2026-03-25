"""
Conversor inteligente de documentos Word a plantillas docxtpl.
Entrenado con los documentos reales de la Alcaldía de Sampués.
"""
import re
import io
from docx import Document

# ── Datos fijos de la alcaldía ─────────────────────────────────────────────
FIJOS = {
    # Municipio y departamento
    "SAMPUÉS – SUCRE": "{{municipio}} – {{departamento}}",
    "SAMPUES –SUCRE": "{{municipio}} –{{departamento}}",
    "SAMPUÉS-SUCRE": "{{municipio}}-{{departamento}}",
    "Sampués-Sucre": "{{municipio}}-{{departamento}}",
    "Sampués - Sucre": "{{municipio}} - {{departamento}}",
    "SAMPUÉS": "{{municipio}}", "Sampués": "{{municipio}}",
    "SAMPUES": "{{municipio}}", "Sampues": "{{municipio}}",
    "SUCRE": "{{departamento}}", "Sucre": "{{departamento}}",

    # Funcionarios
    "CARLOS HUGO MONTOYA ARIAS": "{{nombre_secretario}}",
    "KAROL LAMBRAÑO BUSTAMANTE": "{{nombre_secretario}}",
    "Milta Lambraño Pérez": "{{nombre_elaboro}}",
    "Iván Flórez": "{{nombre_juridico}}",
    "Profesional Universitaria de planeación con funciones asignadas": "{{cargo_secretario}}",
    "Secretaria de Planeación e Infraestructura -Resolución 129-2025": "{{resolucion_cargo}}",
    "Secretario de Planeación e Infraestructura.": "{{cargo_secretario}}",
    "Apoyo a la secretaria de Planeación e infraestructura": "{{cargo_elaboro}}",
    "Profesional de apoyo a la secretaria de planeación": "{{cargo_elaboro}}",
    "Profesional de apoyo a la secretaria de  planeación": "{{cargo_elaboro}}",
    "Asesor Jurídico": "{{cargo_juridico}}",
    "Asesor Jurídica Externo": "{{cargo_juridico}}",

    # Contacto
    "secretariadeplaneacion@sampues-sucre.gov.co": "{{correo_secretaria}}",
    "www.sampues-sucre.gov.co": "{{web_alcaldia}}",
    "283 01 71": "{{telefono_secretaria}}",
    "calle 23 # 19 B-22sector centro": "{{direccion_secretaria}}",
    "705070": "{{codigo_postal}}",

    # Vecinos hardcodeados del auto-cons (datos de ejemplo que no deben quedar)
    "GUSTAVO CARO MERCADO": "{{vecino_1_nombre}}",
    "GLENDA POLO VILLALBA": "{{vecino_2_nombre}}",
    "ANA JULIETA GARCIA ARROYO": "{{vecino_3_nombre}}",
    "C 23C 12 120 LO N°5": "{{vecino_1_direccion}}",
    "C 23C 12 120 LO N° 7": "{{vecino_2_direccion}}",
    "C 23C 13 120 MZ 2 L 5": "{{vecino_3_direccion}}",
    "706700100000001510061000000000": "{{vecino_1_catastro}}",
    "706700100000001510063000000000": "{{vecino_2_catastro}}",
    "706700100000001510033000000000": "{{vecino_3_catastro}}",
    "Barrio el Carmelo": "{{vecino_barrio}}",
}

# ── Patrones con contexto (párrafo completo) → variable específica ─────────
# (palabras_clave_en_parrafo, patron_fecha, variable_resultado)
FECHAS_CONTEXTO = [
    # Vencimiento
    (['vencimiento', 'vence', 'expira', 'vigente hasta'],
     r'\b\d{1,2}\s+de\s+\w+\s+de\s+\d{4}\b', 'fecha_vencimiento'),
    (['vencimiento', 'vence'],
     r'\b\w+\s+\d{1,2}\s+de\s+\d{4}\b', 'fecha_vencimiento'),

    # Radicación
    (['radicó', 'radicacion', 'radicación', 'se recibió', 'recibió oficio', 'se presentó'],
     r'\b\d{1,2}\s+de\s+\w+\s+de\s+\d{4}\b', 'fecha_radicacion'),

    # Notificación
    (['notifique', 'notificación', 'notificacion', 'notifiqué'],
     r'\b\d{1,2}\s+de\s+\w+\s+de\s+\d{4}\b', 'fecha_notificacion'),

    # Expedición (fallback)
    (['expedición', 'expedicion', 'expide', 'expidió', 'expedición:'],
     r'\b\d{1,2}\s+de\s+\w+\s+de\s+\d{4}\b', 'fecha_expedicion'),
    (['expedición', 'expedicion'],
     r'\b\w+\s+\d{1,2}\s+de\s+\d{4}\b', 'fecha_expedicion'),
]

# ── Patrones regex generales (run a run) ──────────────────────────────────
PATRONES = [
    # Radicado municipal (ej: 70670-0-25-0127)
    (r'\b\d{5}-\d{1,2}-\d{2,4}-\d{3,6}\b',          'radicado'),
    # Número de resolución (ej: N° 113, No. 5)
    (r'N[°oº\.]\s*\d{1,4}\b',                         'numero_resolucion'),
    # Cédula (ej: 64.544.148 o 1.234.567)
    (r'\b\d{1,2}[\.\s]?\d{3}[\.\s]?\d{3}\b',         'cedula_solicitante'),
    # Matrícula inmobiliaria (ej: 340-88957)
    (r'\b\d{3}-\d{4,6}\b',                            'matricula_inmobiliaria'),
    # Código catastral (16+ dígitos)
    (r'\b\d{16,30}\b',                                 'codigo_catastral'),
    # Escritura notarial (ej: Escritura No. 1.184)
    (r'[Ee]scritura\s+N[o°\.]+\s*[\d\.]+',            'escritura_predio'),
    # Matrícula profesional (ej: 03-21455 C.P.N.I)
    (r'\b\d{2,6}-\d{4,6}\s*C\.?P\.?[A-Z\.]*\b',      'matricula_prof'),
    # Áreas en hectáreas con m2 (ej: 14 HAS + 5.000,00 M2)
    (r'\d+\s*HAS\s*\+\s*[\d\.,]+\s*M2?',             'area_total_predio'),
    # Áreas solo en hectáreas (ej: 5 HAS)
    (r'\b\d+\s*HAS\b',                                'area_lote'),
    # Fechas escritas — fallback genérico
    (r'\b\d{1,2}\s+de\s+\w+\s+de\s+\d{4}\b',        'fecha_expedicion'),
    (r'\b\w+\s+\d{1,2}\s+de\s+\d{4}\b',              'fecha_expedicion'),
    # Año en contexto (ej: DEL 2025)
    (r'\bDEL\s+20\d{2}\b',                            'anio_resolucion'),
]

def _nombre_unico(base, contadores):
    contadores[base] = contadores.get(base, 0) + 1
    sufijo = f"_{contadores[base]}" if contadores[base] > 1 else ""
    return base + sufijo

def _texto_parrafo(parrafo):
    return ''.join(r.text for r in parrafo.runs)

def _detectar_variable_fecha(texto_parrafo, patron_fecha, contadores):
    """Detecta qué tipo de fecha es según el contexto del párrafo."""
    texto_lower = texto_parrafo.lower()
    for palabras, patron, variable in FECHAS_CONTEXTO:
        if any(p in texto_lower for p in palabras):
            if re.search(patron, texto_parrafo):
                return variable
    return None  # sin contexto claro → usar fallback

def aplicar_fijos(texto):
    for original, variable in FIJOS.items():
        texto = texto.replace(original, variable)
    return texto

def aplicar_patrones_run(texto, contadores, variable_fecha_override=None):
    """Aplica patrones regex al texto de un run."""
    for patron, nombre_base in PATRONES:
        es_fecha = 'fecha' in nombre_base
        def reemplazar(m, nb=nombre_base, es_f=es_fecha, override=variable_fecha_override):
            matched = m.group(0)
            if '{{' in matched:
                return matched
            var = override if (es_f and override) else nb
            return '{{' + _nombre_unico(var, contadores) + '}}'
        texto = re.sub(patron, reemplazar, texto)
    return texto

def procesar_parrafo(parrafo, contadores):
    texto_completo = _texto_parrafo(parrafo)
    if not texto_completo.strip():
        return

    # Detectar tipo de fecha por contexto del párrafo completo
    variable_fecha = _detectar_variable_fecha(texto_completo, None, {})

    for run in parrafo.runs:
        if not run.text.strip():
            continue
        texto = aplicar_fijos(run.text)
        texto = aplicar_patrones_run(texto, contadores, variable_fecha_override=variable_fecha)
        run.text = texto

def convertir_a_plantilla(input_buffer):
    """Recibe un BytesIO con el .docx y devuelve un BytesIO con la plantilla."""
    doc = Document(input_buffer)
    contadores = {}

    for parrafo in doc.paragraphs:
        procesar_parrafo(parrafo, contadores)

    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    procesar_parrafo(parrafo, contadores)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output, list(contadores.keys())
