"""
Conversor inteligente de documentos Word a plantillas docxtpl.
Detecta automáticamente patrones comunes y los reemplaza por {{variables}}.
Uso: python convertir_plantilla.py <ruta_archivo.docx>
     python convertir_plantilla.py  (usa el archivo por defecto)
"""
import re
import sys
import os
from docx import Document

# ── Patrones automáticos (regex → nombre de variable) ─────────────────────
# Se aplican en orden — el primero que coincida gana
PATRONES = [
    # Radicado municipal (ej: 70670-0-25-0127)
    (r'\b\d{5}-\d{1,2}-\d{2,4}-\d{3,6}\b', 'radicado'),

    # Número de resolución (ej: N° 113, No. 5, Nº 22)
    (r'N[°oº\.]\s*\d{1,4}\b', 'numero_resolucion'),

    # Cédula de ciudadanía (ej: 64.544.148 o 64544148)
    (r'\b\d{2}[\.\s]?\d{3}[\.\s]?\d{3}\b', 'cedula_solicitante'),

    # Matrícula inmobiliaria (ej: 340-88957)
    (r'\b\d{3}-\d{4,6}\b', 'matricula_inmobiliaria'),

    # Código catastral (16+ dígitos)
    (r'\b\d{16,30}\b', 'codigo_catastral'),

    # Fechas escritas (ej: 30 de diciembre de 2025, diciembre 30 de 2025)
    (r'\b\d{1,2}\s+de\s+\w+\s+de\s+\d{4}\b', 'fecha_expedicion'),
    (r'\b\w+\s+\d{1,2}\s+de\s+\d{4}\b', 'fecha_expedicion'),

    # Año solo en contexto (ej: DEL 2025, del 2025)
    (r'\bDEL\s+20\d{2}\b', 'anio_resolucion'),

    # Matrícula profesional (ej: 03-21455 C.P.N.I, 12345-CP)
    (r'\b\d{2,6}-\d{4,6}\s*C\.?P\.?[A-Z\.]*\b', 'matricula_prof'),

    # Áreas en hectáreas (ej: 14 HAS, 5 HAS, 9 HAS + 5.000,00 M2)
    (r'\d+\s*HAS\s*\+?\s*[\d\.,]*\s*M2?', 'area_total_predio'),
    (r'\d+\s*HAS\b', 'area_lote'),

    # Áreas en m2 (ej: 120,50 M2, 1.200 m2)
    (r'\b[\d\.,]+\s*[Mm]2\b', 'area_m2'),
]

# ── Reemplazos fijos conocidos (texto exacto → variable) ──────────────────
# Agrega aquí valores específicos de tu municipio que siempre serán iguales
FIJOS = {
    "SAMPUÉS": "{{municipio}}",
    "Sampués": "{{municipio}}",
    "SAMPUES": "{{municipio}}",
    "Sampues": "{{municipio}}",
    "SUCRE": "{{departamento}}",
    "Sucre": "{{departamento}}",
    "CARLOS HUGO MONTOYA ARIAS": "{{nombre_secretario}}",
    "Milta Lambraño Pérez": "{{nombre_elaboro}}",
    "Iván Flórez": "{{nombre_juridico}}",
    "secretariadeplaneacion@sampues-sucre.gov.co": "{{correo_secretaria}}",
    "www.sampues-sucre.gov.co": "{{web_alcaldia}}",
    "283 01 71": "{{telefono_secretaria}}",
    "calle 23 # 19 B-22sector centro": "{{direccion_secretaria}}",
    "705070": "{{codigo_postal}}",
}

# Contador para variables repetidas (area_lote_1, area_lote_2, etc.)
_contadores = {}

def nombre_variable_unico(base):
    if base not in _contadores:
        _contadores[base] = 0
        return base
    _contadores[base] += 1
    return f"{base}_{_contadores[base]}"

def aplicar_patrones(texto):
    """Aplica reemplazos fijos y luego patrones regex."""
    # 1. Fijos primero
    for original, variable in FIJOS.items():
        texto = texto.replace(original, variable)

    # 2. Regex — solo si no está ya dentro de {{ }}
    for patron, nombre_base in PATRONES:
        def reemplazar(m):
            matched = m.group(0)
            # No reemplazar si ya es una variable
            if '{{' in matched:
                return matched
            var = nombre_variable_unico(nombre_base)
            return '{{' + var + '}}'
        texto = re.sub(patron, reemplazar, texto)

    return texto

def procesar_run(run):
    if not run.text.strip():
        return
    nuevo = aplicar_patrones(run.text)
    if nuevo != run.text:
        run.text = nuevo

def procesar_documento(input_path, output_path):
    doc = Document(input_path)

    # Párrafos del cuerpo
    for parrafo in doc.paragraphs:
        for run in parrafo.runs:
            procesar_run(run)

    # Tablas
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    for run in parrafo.runs:
                        procesar_run(run)

    doc.save(output_path)
    print(f"✅ Plantilla guardada en: {output_path}")
    print(f"   Variables detectadas: {list(_contadores.keys()) + list(set(FIJOS.values()))}")

if __name__ == "__main__":
    if len(sys.argv) > 1:
        input_file = sys.argv[1]
    else:
        input_file = r"C:\Users\Jesus\Downloads\subdivision-cons.docx"

    if not os.path.exists(input_file):
        print(f"❌ No se encontró: {input_file}")
        print("Uso: python convertir_plantilla.py <ruta_al_archivo.docx>")
        sys.exit(1)

    base = os.path.splitext(input_file)[0]
    output_file = base + "-plantilla.docx"
    procesar_documento(input_file, output_file)
